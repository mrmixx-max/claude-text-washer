from pathlib import Path
from scripts.file_washer import detect_format, extract_file_text, expand_inputs, output_path_for
from scripts.formats.documents import extract_text

import pymupdf, io
from docx import Document

# Build fixtures
tmp = Path("C:/Users/webma/Downloads/claude-text-washer/_fixtures")
tmp.mkdir(exist_ok=True)

docx_path = tmp / "sample.docx"
doc = Document(); doc.add_paragraph("Hello docx world"); doc.add_paragraph("Second paragraph here")
doc.save(str(docx_path))

pdf_path = tmp / "sample.pdf"
mdoc = pymupdf.open(); page = mdoc.new_page(); page.insert_text((72, 200), "Hello PDF text world")
pdf_path.write_bytes(mdoc.write()); mdoc.close()

md_path = tmp / "sample.md"
md_path.write_text("# Title\n\nSome **markdown** text.\n\n- list item one\n- list item two\n")

html_path = tmp / "sample.html"
html_path.write_text("<!doctype html><html><head><meta name='author' content='x'><title>T</title></head><body><h1>Hi</h1><p>Hello HTML paragraph.</p></body></html>")

txt_path = tmp / "sample.txt"
txt_path.write_text("Plain text file.\nSecond line.")

for p in [docx_path, pdf_path, md_path, html_path, txt_path]:
    fmt = detect_format(p)
    text = extract_file_text(p)
    print(f"{p.name:12s} fmt={fmt:6s} chars={len(text):4d} text={text[:40]!r}")

# Test force format override
print("pdf forced as txt:", repr(extract_file_text(pdf_path, fmt="txt")[:30]))

# Test expand_inputs with glob + directory
files = expand_inputs([str(pdf_path), str(tmp / "*.md"), str(tmp / "*.docx")])
print("expanded:", [f.name for f in files])
assert len(files) == 3, files

# Test directory walk
files2 = expand_inputs([str(tmp)], recursive=False)
print("dir walk:", sorted(f.name for f in files2))

# Test output path
print("output single:", output_path_for(docx_path, None, None))
print("output outdir:", output_path_for(docx_path, None, str(tmp / "out")))
print("ALL file_washer extraction OK")
